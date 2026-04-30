"""
QMT API文档爬取工具
"""
import requests
from bs4 import BeautifulSoup
from docx import Document
try:
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    pdfkit = None
    PDFKIT_AVAILABLE = False
import os
import time
from urllib.parse import urljoin, urlparse

class QMTAPIScraper:
    """QMT API文档爬取器"""
    
    def __init__(self, base_url="https://dict.thinktrader.net"):
        self.base_url = base_url
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
    
    def get_page_content(self, url):
        """获取页面内容"""
        try:
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
            response.encoding = 'utf-8'
            return response.text
        except Exception as e:
            print(f"获取页面失败 {url}: {e}")
            return None
    
    def parse_main_page(self):
        """解析主页面，获取所有API分类链接"""
        main_url = self.base_url
        content = self.get_page_content(main_url)
        if not content:
            return []
        
        soup = BeautifulSoup(content, 'html.parser')
        links = []
        
        # 查找所有API分类链接
        for link in soup.find_all('a', href=True):
            href = link['href']
            if '/dictionary/' in href and href.endswith('.html'):
                full_url = urljoin(self.base_url, href)
                links.append({
                    'url': full_url,
                    'title': link.get_text(strip=True) or href.split('/')[-1].replace('.html', '')
                })
        
        return links
    
    def scrape_api_content(self, url, title):
        """爬取单个API页面内容"""
        content = self.get_page_content(url)
        if not content:
            return ""
        
        soup = BeautifulSoup(content, 'html.parser')
        
        # 移除不必要的元素
        for element in soup.find_all(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()
        
        # 创建文档内容
        doc_content = f"<h1>{title}</h1>\n"
        
        # 提取主要内容
        main_content = soup.find('body')
        if main_content:
            # 移除侧边栏等不需要的内容
            for sidebar in main_content.find_all(class_=lambda x: x and ('sidebar' in x or 'nav' in x)):
                sidebar.decompose()
            
            doc_content += str(main_content)
        
        return doc_content
    
    def create_word_document(self, content_dict, output_path):
        """创建Word文档"""
        doc = Document()
        doc.add_heading('QMT API文档', 0)
        
        # 添加目录页
        doc.add_paragraph('目录', style='Heading 1')
        for title in content_dict.keys():
            doc.add_paragraph(title, style='List Bullet')
        
        # 添加各章节内容
        for title, content in content_dict.items():
            doc.add_heading(title, level=1)
            # 简单处理HTML内容转为Word内容
            soup = BeautifulSoup(content, 'html.parser')
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                level = int(element.name[1]) if element.name[1].isdigit() else 1
                doc.add_heading(element.get_text(strip=True), level=min(level, 6))
                element.decompose()
            
            # 处理代码块
            for pre in soup.find_all('pre'):
                code_text = pre.get_text()
                if code_text.strip():
                    doc.add_paragraph(code_text, style='Normal')
                pre.decompose()
            
            # 处理表格
            for table in soup.find_all('table'):
                # 简单处理表格，转换为文本
                table_text = ""
                for row in table.find_all('tr'):
                    cells = row.find_all(['td', 'th'])
                    row_text = " | ".join([cell.get_text(strip=True) for cell in cells])
                    table_text += row_text + "\n"
                if table_text.strip():
                    doc.add_paragraph(table_text)
                table.decompose()
            
            # 处理段落
            for p in soup.find_all('p'):
                text = p.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
        
        doc.save(output_path)
        print(f"Word文档已保存到: {output_path}")
    
    def create_pdf_document(self, html_content, output_path):
        """创建PDF文档"""
        if not PDFKIT_AVAILABLE:
            print("PDF功能不可用: 未安装pdfkit或wkhtmltopdf")
            return
            
        try:
            # 添加基本的HTML结构
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>QMT API文档</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; }}
                    h1, h2, h3 {{ color: #333; }}
                    pre {{ background-color: #f5f5f5; padding: 10px; overflow-x: auto; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                </style>
            </head>
            <body>
                <h1>QMT API文档</h1>
                {html_content}
            </body>
            </html>
            """
            
            # 配置pdfkit选项
            options = {
                'page-size': 'A4',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None
            }
            
            pdfkit.from_string(full_html, output_path, options=options)
            print(f"PDF文档已保存到: {output_path}")
        except Exception as e:
            print(f"创建PDF文档失败: {e}")
    
    def scrape_all_apis(self):
        """爬取所有API文档"""
        print("开始爬取QMT API文档...")
        
        # 获取所有API分类链接
        api_links = self.parse_main_page()
        print(f"找到 {len(api_links)} 个API分类")
        
        # 存储所有内容
        content_dict = {}
        all_html_content = ""
        
        # 爬取每个API页面
        for i, link_info in enumerate(api_links):
            url = link_info['url']
            title = link_info['title']
            print(f"正在爬取 ({i+1}/{len(api_links)}): {title}")
            
            content = self.scrape_api_content(url, title)
            if content:
                content_dict[title] = content
                all_html_content += f"<div><h1>{title}</h1>{content}</div>"
            
            # 添加延时，避免请求过于频繁
            time.sleep(1)
        
        # 创建输出目录
        output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'api_docs')
        os.makedirs(output_dir, exist_ok=True)
        
        # 生成Word文档
        word_path = os.path.join(output_dir, 'QMT_API文档.docx')
        self.create_word_document(content_dict, word_path)
        
        # 生成PDF文档
        pdf_path = os.path.join(output_dir, 'QMT_API文档.pdf')
        self.create_pdf_document(all_html_content, pdf_path)
        
        # 如果PDF生成失败，创建HTML文档作为替代
        if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
            html_path = os.path.join(output_dir, 'QMT_API文档.html')
            try:
                # 添加基本的HTML结构
                full_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <title>QMT API文档</title>
                    <style>
                        body {{ font-family: Arial, sans-serif; margin: 40px; }}
                        h1, h2, h3 {{ color: #333; }}
                        pre {{ background-color: #f5f5f5; padding: 10px; overflow-x: auto; }}
                        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
                        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                        th {{ background-color: #f2f2f2; }}
                    </style>
                </head>
                <body>
                    <h1>QMT API文档</h1>
                    {all_html_content}
                </body>
                </html>
                """
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(full_html)
                print(f"HTML文档已保存到: {html_path}")
            except Exception as e:
                print(f"创建HTML文档失败: {e}")
        
        print("API文档爬取完成!")
        return word_path, pdf_path

def main():
    """主函数"""
    scraper = QMTAPIScraper()
    word_path, pdf_path = scraper.scrape_all_apis()
    print(f"文档已保存到:")
    print(f"  Word: {word_path}")
    print(f"  PDF: {pdf_path}")

if __name__ == "__main__":
    main()